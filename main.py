from openai import OpenAI
import crendentials as c

client = OpenAI(
    api_key=c.MyAPI
)

filename = "insert_filename_here"
extension = "mp3"
dodebug = False

systemPrompts = {
    "abstractSystemPrompt": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points.",
    "keyPointsSystemPrompt": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about.",
    "actionItemsSystemPrompt": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely.",
    "sentimentSystemPrompt": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."
}

from docx import Document

def debug(msg):
    if dodebug:
        print()
        print(msg)

def transcribe_audio(audio_file_path):
    debug("transcribe start")
    with open(audio_file_path, 'rb') as audio_file:
        transcription = client.audio.transcriptions.create(
            model="whisper-1",
            file=audio_file
        )
    return transcription.text

def meeting_minutes(transcription):
    debug('running meeting minutes')
    abstract_summary = extraction(0, systemPrompts["abstractSystemPrompt"], transcription)
    key_points = extraction(0, systemPrompts["keyPointsSystemPrompt"], transcription)
    action_items = extraction(0, systemPrompts["actionItemsSystemPrompt"], transcription)
    sentiment = extraction(0, systemPrompts["sentimentSystemPrompt"], transcription)
    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }

def extraction(temp, prompt, transcription):
    debug('abstract_summary_extraction')
    response = client.chat.completions.create(
        model="gpt-4o",
        temperature=temp,
        messages=[
            {
                "role": "system",
                "content": prompt
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def save_as_docx_extraction_data(minutes, filename):
    debug("save_as_docx")
    doc = Document()
    doc.add_paragraph(minutes)
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)

def save_as_docx_transcription(minutes, filename):
    debug("save_as_docx")
    doc = Document()
    doc.add_paragraph(minutes)
    doc.save(filename)

def main():
    audio_file_path = f"audio/{filename}.{extension}"
    transcription = transcribe_audio(audio_file_path)
    debug(transcription)
    save_as_docx_transcription(transcription, f"audio/{filename}-transcript")
    save_as_docx_extraction_data(transcription, f"audio/{filename}-notes.docx")

if __name__ == "__main__":
    main()